import os
import io
from typing import Dict, Any, Optional
from PIL import Image
from backend.app.utils.logger import logger

def extract_text_from_pdf(file_path: str) -> str:
    text_content = []
    
    # 1. Try PyMuPDF (fitz)
    try:
        import fitz
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text and text.strip():
                text_content.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
        doc.close()
        if text_content:
            return "\n\n".join(text_content)
    except Exception as e:
        logger.warning(f"PyMuPDF failed on {file_path}: {e}")

    # 2. Try pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    text_content.append(f"--- Page {i + 1} ---\n{text.strip()}")
        if text_content:
            return "\n\n".join(text_content)
    except Exception as e:
        logger.warning(f"pdfplumber failed on {file_path}: {e}")

    # 3. Try pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                text_content.append(f"--- Page {i + 1} ---\n{text.strip()}")
        if text_content:
            return "\n\n".join(text_content)
    except Exception as e:
        logger.warning(f"pypdf failed on {file_path}: {e}")

    return "\n\n".join(text_content) if text_content else "No extractable text found in PDF."

def extract_text_from_image(file_path: str) -> str:
    try:
        import pytesseract
        image = Image.open(file_path)
        ocr_text = pytesseract.image_to_string(image)
        if ocr_text and ocr_text.strip():
            return ocr_text.strip()
    except Exception as e:
        logger.warning(f"Tesseract OCR not configured or failed for {file_path}: {e}")
    
    # Return placeholder indication if local OCR is unavailable so Gemini multimodal / mock can process
    return f"[Image file processed: {os.path.basename(file_path)}]"

def extract_text_from_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"docx extraction failed for {file_path}: {e}")
        return ""

def process_uploaded_document(file_path: str) -> Dict[str, Any]:
    ext = os.path.splitext(file_path)[1].lower()
    raw_text = ""
    file_size = 0
    try:
        file_size = os.path.getsize(file_path)
    except Exception:
        pass

    if ext in [".pdf"]:
        raw_text = extract_text_from_pdf(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
        raw_text = extract_text_from_image(file_path)
    elif ext in [".docx", ".doc"]:
        raw_text = extract_text_from_docx(file_path)
    elif ext in [".txt", ".csv", ".json", ".md"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        except Exception as e:
            logger.warning(f"Text reading failed: {e}")
    else:
        raw_text = f"Binary/Unsupported format file: {os.path.basename(file_path)}"

    return {
        "file_path": file_path,
        "file_name": os.path.basename(file_path),
        "file_type": ext.lstrip("."),
        "file_size": file_size,
        "raw_text": raw_text
    }
